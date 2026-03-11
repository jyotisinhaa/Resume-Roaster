"""
PDF / DOCX / TXT extraction utilities for Resume Roaster.
"""

import fitz  # PyMuPDF
from docx import Document
import io
import logging

# ── Logging setup ─────────────────────────────────────────────────────────────
import os as _os
_LOG_DIR = _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))
_LOG_FILE = _os.path.join(_LOG_DIR, "resume_roaster.log")

logger = logging.getLogger("ResumeRoaster")
logger.setLevel(logging.DEBUG)
if not logger.handlers:
    # File handler — writes all logs to resume_roaster.log
    fh = logging.FileHandler(_LOG_FILE, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fmt = logging.Formatter(
        "%(asctime)s | %(name)s | %(levelname)-7s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )
    fh.setFormatter(fmt)
    logger.addHandler(fh)
    # Also print to console
    ch = logging.StreamHandler()
    ch.setLevel(logging.DEBUG)
    ch.setFormatter(fmt)
    logger.addHandler(ch)

# ── Constants ─────────────────────────────────────────────────────────────────
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024  # 10 MB


def validate_file(file_bytes: bytes, file_name: str) -> None:
    """
    Validate uploaded file:
      1. Check file extension is supported (.pdf, .docx, .txt)
      2. Check file size is within the limit (10 MB)
    Raises ValueError with a descriptive message on failure.
    """
    import os
    _, ext = os.path.splitext(file_name.lower())
    size_mb = len(file_bytes) / (1024 * 1024)

    logger.info(f"📄 Validating file: {file_name}")
    logger.info(f"   Extension: {ext}")
    logger.info(f"   Size: {size_mb:.2f} MB ({len(file_bytes):,} bytes)")

    if ext not in ALLOWED_EXTENSIONS:
        logger.warning(f"❌ Rejected: unsupported extension '{ext}'")
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Please upload one of: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        logger.warning(f"❌ Rejected: file too large ({size_mb:.1f} MB)")
        raise ValueError(
            f"File too large ({size_mb:.1f} MB). "
            f"Maximum allowed size is {MAX_FILE_SIZE_MB} MB."
        )

    logger.info(f"✅ File validation passed")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF stored as bytes."""
    logger.info("📖 Extracting text from PDF...")
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts: list[str] = []
    for i, page in enumerate(doc):
        page_text = page.get_text()
        text_parts.append(page_text)
        logger.debug(f"   Page {i+1}: {len(page_text)} characters")
    page_count = len(doc)
    doc.close()
    full_text = "\n".join(text_parts).strip()
    logger.info(f"✅ PDF extracted: {page_count} pages, {len(full_text):,} total characters")
    logger.info(f"📝 FULL EXTRACTED TEXT:\n{'─'*60}\n{full_text}\n{'─'*60}")
    return full_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX stored as bytes."""
    logger.info("📖 Extracting text from DOCX...")
    doc = Document(io.BytesIO(file_bytes))
    text_parts: list[str] = []
    para_count = 0
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
            para_count += 1
    logger.debug(f"   Paragraphs extracted: {para_count}")
    # Also extract text from tables
    table_count = 0
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
                table_count += 1
    logger.debug(f"   Table rows extracted: {table_count}")
    full_text = "\n".join(text_parts).strip()
    logger.info(f"✅ DOCX extracted: {para_count} paragraphs, {table_count} table rows, {len(full_text):,} total characters")
    logger.info(f"📝 FULL EXTRACTED TEXT:\n{'─'*60}\n{full_text}\n{'─'*60}")
    return full_text


def extract_text(file_bytes: bytes, file_name: str) -> str:
    """
    Validate then route to the correct parser based on file extension.
    Supports .pdf, .docx, and .txt files.
    """
    logger.info(f"{'='*60}")
    logger.info(f"🔥 Processing file: {file_name}")
    logger.info(f"{'='*60}")

    # Step 1: Validate
    validate_file(file_bytes, file_name)

    # Step 2: Extract
    name_lower = file_name.lower()
    if name_lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    elif name_lower.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="replace").strip()
        logger.info(f"✅ TXT extracted: {len(text):,} characters")
        logger.info(f"📝 FULL EXTRACTED TEXT:\n{'─'*60}\n{text}\n{'─'*60}")
    else:
        raise ValueError(f"Unsupported file type: {file_name}")

    logger.info(f"{'='*60}")
    logger.info(f"✅ Extraction complete for {file_name}")
    logger.info(f"   Total text length: {len(text):,} characters")
    logger.info(f"{'='*60}")
    return text
