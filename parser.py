"""
PDF / DOCX / TXT extraction utilities for Resume Roaster.
"""

import fitz  # PyMuPDF
from docx import Document
import io

# ── Constants ─────────────────────────────────────────────────────────────────
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024  # 10 MB


def validate_file(file_bytes: bytes, file_name: str) -> None:
    """
    Validate uploaded file:
      1. Check file extension is supported (.pdf, .docx, .txt)
      2. Check file size is within the limit (10 MB)
    Raises ValueError with a descriptive message on failure.
    """
    import os
    _, ext = os.path.splitext(file_name.lower())

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Please upload one of: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    size_mb = len(file_bytes) / (1024 * 1024)
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"File too large ({size_mb:.1f} MB). "
            f"Maximum allowed size is {MAX_FILE_SIZE_MB} MB."
        )


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF stored as bytes."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts: list[str] = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    print(f"Extracted text from PDF: {len(text_parts)} pages, {sum(len(p) for p in text_parts)} characters")
    return "\n".join(text_parts).strip()
    print(f"text_parts: {text_parts[:3]}...")  # Debug: Show first 3 parts


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX stored as bytes."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts).strip()


def extract_text(file_bytes: bytes, file_name: str) -> str:
    """
    Validate then route to the correct parser based on file extension.
    Supports .pdf, .docx, and .txt files.
    """
    # Step 1: Validate
    validate_file(file_bytes, file_name)

    # Step 2: Extract
    name_lower = file_name.lower()
    if name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif name_lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace").strip()
    else:
        raise ValueError(f"Unsupported file type: {file_name}")
